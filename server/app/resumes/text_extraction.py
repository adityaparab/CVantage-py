"""Text extraction service for resume files (issue #45).

Extracts plain text from PDF, DOCX, and DOC file formats.
"""

from __future__ import annotations

import io
import re
from typing import Literal

import docx
import pypdf

MAX_TEXT_CHARS = 200_000


class ExtractionError(Exception):
    """Base error for all extraction failures."""

    def __init__(self, message: str, reason: str) -> None:
        self.message = message
        self.reason = reason
        super().__init__(message)


class EncryptedFileError(ExtractionError):
    """The PDF is password-protected and cannot be extracted."""

    def __init__(self) -> None:
        super().__init__("File is encrypted or password-protected", "encrypted")


class CorruptFileError(ExtractionError):
    """The file is corrupt or not a valid document."""

    def __init__(self, detail: str = "") -> None:
        msg = "File is corrupt or not a valid document"
        if detail:
            msg += f": {detail}"
        super().__init__(msg, "corrupt")


class EmptyTextError(ExtractionError):
    """The file contains no extractable text."""

    def __init__(self) -> None:
        super().__init__("No text could be extracted from the file", "empty")


ExtractionResult = tuple[str, Literal["pdf", "docx", "doc"]]


def _normalize(text: str) -> str:
    """Normalize whitespace and encoding, cap at MAX_TEXT_CHARS."""
    text = text.encode("utf-8", errors="replace").decode("utf-8")
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS]
    return text


async def extract_pdf(data: bytes) -> ExtractionResult:
    """Extract text from a PDF document using pypdf."""
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
    except pypdf.errors.PdfReadError as e:
        raise CorruptFileError(str(e)) from e

    if reader.is_encrypted:
        raise EncryptedFileError()

    try:
        pages = [page.extract_text() for page in reader.pages]
    except Exception as e:
        raise CorruptFileError(str(e)) from e

    text = "\n".join(pages)
    text = _normalize(text)

    if not text:
        raise EmptyTextError()

    return text, "pdf"


async def extract_docx(data: bytes) -> ExtractionResult:
    """Extract text from a DOCX document using python-docx."""
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as e:
        raise CorruptFileError(str(e)) from e

    paragraphs = [p.text for p in document.paragraphs]
    text = "\n".join(paragraphs)
    text = _normalize(text)

    if not text:
        raise EmptyTextError()

    return text, "docx"


async def extract_doc(data: bytes) -> ExtractionResult:
    """Extract text from a legacy DOC document using python-docx (limited support).

    python-docx has basic support for .doc files via the OLE2 container.
    If that fails, we try a raw text extraction fallback.
    """
    try:
        document = docx.Document(io.BytesIO(data))
        paragraphs = [p.text for p in document.paragraphs]
        text = "\n".join(paragraphs)
        text = _normalize(text)
        if text:
            return text, "doc"
    except Exception:
        pass

    # Fallback: try extracting readable text from the binary
    try:
        raw = data.decode("latin-1")
        # Filter printable ASCII/Unicode sequences
        words = re.findall(r"[\x20-\x7E\u00A0-\u024F]{3,}", raw)
        text = " ".join(words)
        text = _normalize(text)
        if text:
            return text, "doc"
    except Exception as e:
        raise CorruptFileError(str(e)) from e

    raise EmptyTextError()


async def extract_text(data: bytes, ext: str) -> ExtractionResult:
    """Extract text from a resume file based on its extension.

    Args:
        data: Raw file contents.
        ext: File extension ('.pdf', '.docx', '.doc').

    Returns:
        Tuple of (extracted_text, format_name).

    Raises:
        EncryptedFileError: PDF is password-protected.
        CorruptFileError: File is corrupt.
        EmptyTextError: No text could be extracted.
    """
    ext_lower = ext.lower()
    if ext_lower == ".pdf":
        return await extract_pdf(data)
    elif ext_lower == ".docx":
        return await extract_docx(data)
    elif ext_lower == ".doc":
        return await extract_doc(data)
    else:
        raise CorruptFileError(f"Unsupported format: {ext}")

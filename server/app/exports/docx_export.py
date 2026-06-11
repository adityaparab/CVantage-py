"""DOCX resume export (issue #90) via python-docx.

Maps every json-resume section into a Word document. Input is the by-alias
json-resume dict (camelCase keys: startDate, etc.).
"""

from __future__ import annotations

import io
from collections.abc import Callable
from typing import Any

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _date_range(item: dict[str, Any]) -> str:
    start = item.get("startDate") or item.get("date") or item.get("releaseDate") or ""
    end = item.get("endDate") or ""
    if start and end:
        return f"{start} – {end}"
    return str(start or end or "")


def render_docx(name: str, json_resume: dict[str, Any]) -> bytes:
    document = docx.Document()
    basics = json_resume.get("basics") or {}

    title = document.add_heading(basics.get("name") or name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if basics.get("label"):
        sub = document.add_paragraph(basics["label"])
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_bits = [basics.get("email"), basics.get("phone"), basics.get("url")]
    contact = " · ".join(b for b in contact_bits if b)
    if contact:
        para = document.add_paragraph(contact)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if basics.get("summary"):
        document.add_heading("Summary", level=1)
        document.add_paragraph(basics["summary"])

    def list_section(
        heading: str, items: list[dict[str, Any]], render: Callable[[dict[str, Any]], None]
    ) -> None:
        if not items:
            return
        document.add_heading(heading, level=1)
        for item in items:
            render(item)

    def render_work(item: dict[str, Any]) -> None:
        head = document.add_paragraph()
        run = head.add_run(f"{item.get('position', '')} — {item.get('name', '')}".strip(" —"))
        run.bold = True
        run.font.size = Pt(11)
        dates = _date_range(item)
        if dates:
            document.add_paragraph(dates)
        if item.get("summary"):
            document.add_paragraph(item["summary"])
        for highlight in item.get("highlights") or []:
            document.add_paragraph(highlight, style="List Bullet")

    list_section("Experience", json_resume.get("work") or [], render_work)
    list_section("Volunteer", json_resume.get("volunteer") or [], render_work)

    def render_education(item: dict[str, Any]) -> None:
        head = document.add_paragraph()
        run = head.add_run(f"{item.get('institution', '')} — {item.get('area', '')}".strip(" —"))
        run.bold = True
        if item.get("studyType"):
            document.add_paragraph(item["studyType"])
        dates = _date_range(item)
        if dates:
            document.add_paragraph(dates)

    list_section("Education", json_resume.get("education") or [], render_education)

    def render_project(item: dict[str, Any]) -> None:
        head = document.add_paragraph()
        run = head.add_run(item.get("name", ""))
        run.bold = True
        if item.get("description"):
            document.add_paragraph(item["description"])
        for highlight in item.get("highlights") or []:
            document.add_paragraph(highlight, style="List Bullet")

    list_section("Projects", json_resume.get("projects") or [], render_project)

    skills = json_resume.get("skills") or []
    if skills:
        document.add_heading("Skills", level=1)
        document.add_paragraph(", ".join(s.get("name", "") for s in skills if s.get("name")))

    awards = json_resume.get("awards") or []
    if awards:
        document.add_heading("Awards", level=1)
        for award in awards:
            document.add_paragraph(
                f"{award.get('title', '')} — {award.get('awarder', '')}".strip(" —")
            )

    certificates = json_resume.get("certificates") or []
    if certificates:
        document.add_heading("Certificates", level=1)
        for cert in certificates:
            document.add_paragraph(f"{cert.get('name', '')} — {cert.get('issuer', '')}".strip(" —"))

    languages = json_resume.get("languages") or []
    if languages:
        document.add_heading("Languages", level=1)
        document.add_paragraph(
            ", ".join(
                f"{lang.get('language', '')} ({lang.get('fluency', '')})".strip(" ()")
                for lang in languages
            )
        )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
